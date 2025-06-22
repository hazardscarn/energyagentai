"""
Marketing Content Creation Tools
File: marketing_agent/marketing_content_tools.py
COMPLETE UPDATED FILE - Creates Word documents instead of ZIP files
"""

import os
import json
import uuid
from datetime import datetime
from typing import Optional, Dict, List
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Global instances
_content_manager = None
_output_base_dir = "marketing_outputs"

def initialize_content_manager():
    """Initialize Content Manager (call once)"""
    global _content_manager
    if _content_manager is None:
        _content_manager = MarketingContentManager(_output_base_dir)

class MarketingContentManager:
    """Manages marketing content creation and file operations"""
    
    def __init__(self, base_output_dir: str):
        self.base_output_dir = base_output_dir
        self.sessions = {}  # Track user sessions and their content
        
        # Create base directories
        os.makedirs(base_output_dir, exist_ok=True)
        os.makedirs(f"{base_output_dir}/downloads", exist_ok=True)
    
    def create_session_dir(self, session_id: str) -> str:
        """Create directory for a user session"""
        session_dir = f"{self.base_output_dir}/{session_id}"
        os.makedirs(session_dir, exist_ok=True)
        os.makedirs(f"{session_dir}/images", exist_ok=True)
        os.makedirs(f"{session_dir}/content", exist_ok=True)
        return session_dir
    
    def save_marketing_content(self, content: str, content_type: str, 
                             session_id: str, filename: Optional[str] = None) -> Dict:
        """Save marketing content to file"""
        try:
            session_dir = self.create_session_dir(session_id)
            
            if filename is None:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"{content_type}_{timestamp}.txt"
            
            # Ensure .txt extension for text content
            if not filename.endswith('.txt'):
                filename += '.txt'
            
            file_path = f"{session_dir}/content/{filename}"
            
            # Save content
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            # Track in session
            if session_id not in self.sessions:
                self.sessions[session_id] = {'content_files': [], 'image_files': []}
            
            self.sessions[session_id]['content_files'].append({
                'filename': filename,
                'path': file_path,
                'type': content_type,
                'created': datetime.now().isoformat()
            })
            
            return {
                'success': True,
                'filename': filename,
                'path': file_path,
                'size_kb': round(os.path.getsize(file_path) / 1024, 2)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
    
    def create_word_document(self, session_id: str) -> Dict:
        """Create downloadable Word document with content and embedded images"""
        try:
            if session_id not in self.sessions:
                return {
                    'success': False,
                    'error': 'No content found for this session'
                }
            
            session_data = self.sessions[session_id]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            doc_filename = f"marketing_campaign_{session_id}_{timestamp}.docx"
            doc_path = f"{self.base_output_dir}/downloads/{doc_filename}"
            
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Marketing Campaign Materials', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add generation info
            info_para = doc.add_paragraph()
            info_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y at %H:%M")}').bold = True
            info_para.add_run(f'\nSession ID: {session_id}')
            
            # Process content files
            content_files = session_data.get('content_files', [])
            image_files = session_data.get('image_files', [])
            
            # Group images by type for easier matching
            images_by_type = {}
            for img in image_files:
                img_type = img['type']
                if img_type not in images_by_type:
                    images_by_type[img_type] = []
                images_by_type[img_type].append(img)
            
            # Process each content file
            for i, content_file in enumerate(content_files):
                if i > 0:
                    doc.add_page_break()
                
                # Add content section header
                content_type = content_file['type'].replace('_', ' ').title()
                doc.add_heading(f'{content_type}', 1)
                
                # Read and add content
                try:
                    with open(content_file['path'], 'r', encoding='utf-8') as f:
                        content = f.read()
                    
                    # Format content nicely
                    self._format_content_in_doc(doc, content, content_type)
                    
                except Exception as e:
                    doc.add_paragraph(f"Error reading content: {str(e)}")
                
                # Add related images
                related_images = self._find_related_images(content_file['type'], images_by_type)
                if related_images:
                    doc.add_heading('Related Images', 2)
                    for img in related_images[:2]:  # Limit to 2 images per content section
                        try:
                            if os.path.exists(img['path']):
                                doc.add_paragraph()
                                
                                # Add image with proper sizing
                                doc.add_picture(img['path'], width=Inches(5.5))
                                
                                # Add image caption
                                caption = doc.add_paragraph()
                                caption.add_run(f"Image: {img['filename']}").italic = True
                                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            doc.add_paragraph(f"Error embedding image {img['filename']}: {str(e)}")
            
            # Add any remaining images not matched to content
            unmatched_images = self._get_unmatched_images(content_files, image_files)
            if unmatched_images:
                doc.add_page_break()
                doc.add_heading('Additional Images', 1)
                for img in unmatched_images:
                    try:
                        if os.path.exists(img['path']):
                            doc.add_paragraph()
                            doc.add_picture(img['path'], width=Inches(5.5))
                            caption = doc.add_paragraph()
                            caption.add_run(f"Image: {img['filename']}").italic = True
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        doc.add_paragraph(f"Error embedding image {img['filename']}: {str(e)}")
            
            # Add usage instructions
            doc.add_page_break()
            doc.add_heading('Usage Instructions', 1)
            instructions = [
                "This document contains your complete marketing campaign materials.",
                "Copy and paste the content into your preferred email marketing platform.",
                "Use the embedded images in your campaigns by right-clicking and saving them.",
                "Customize the content as needed for your specific brand voice and requirements.",
                "For email campaigns, test the formatting in your email platform before sending."
            ]
            
            for instruction in instructions:
                p = doc.add_paragraph()
                p.add_run('• ').bold = True
                p.add_run(instruction)
            
            # Save document
            doc.save(doc_path)
            
            # Get file size
            file_size_mb = round(os.path.getsize(doc_path) / (1024 * 1024), 2)
            
            return {
                'success': True,
                'doc_filename': doc_filename,
                'doc_path': doc_path,
                'download_url': f"marketing_outputs/downloads/{doc_filename}",
                'size_mb': file_size_mb,
                'content_count': len(content_files),
                'image_count': len(image_files)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
    
    def _format_content_in_doc(self, doc, content: str, content_type: str):
        """Format content nicely in Word document"""
        lines = content.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_paragraph:
                    current_paragraph = None
                continue
            
            # Check if line looks like a heading
            if (line.isupper() and len(line) < 100) or line.startswith('#'):
                heading_text = line.replace('#', '').strip()
                doc.add_heading(heading_text, 3)
                current_paragraph = None
            # Check if line looks like a subject line
            elif 'subject:' in line.lower():
                p = doc.add_paragraph()
                p.add_run('Subject: ').bold = True
                p.add_run(line.split(':', 1)[1].strip())
                current_paragraph = None
            else:
                # Regular content
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                else:
                    current_paragraph.add_run('\n')
                current_paragraph.add_run(line)
    
    def _find_related_images(self, content_type: str, images_by_type: dict) -> List:
        """Find images related to a content type"""
        # Direct match first
        if content_type in images_by_type:
            return images_by_type[content_type]
        
        # Try to match based on content type
        content_lower = content_type.lower()
        related = []
        
        for img_type, imgs in images_by_type.items():
            img_type_lower = img_type.lower()
            
            # Match email content with email images
            if 'email' in content_lower and 'email' in img_type_lower:
                related.extend(imgs)
            # Match social content with social images  
            elif 'social' in content_lower and 'social' in img_type_lower:
                related.extend(imgs)
            # Match promotional content with promotional images
            elif 'promotional' in content_lower and 'promotional' in img_type_lower:
                related.extend(imgs)
            # General marketing match
            elif 'marketing' in img_type_lower:
                related.extend(imgs)
        
        return related
    
    def _get_unmatched_images(self, content_files: List, image_files: List) -> List:
        """Get images that weren't matched to any content"""
        used_types = set()
        for content in content_files:
            content_type = content['type'].lower()
            for img in image_files:
                if (content_type in img['type'].lower() or 
                    img['type'].lower() in content_type or
                    'email' in content_type and 'email' in img['type'].lower() or
                    'social' in content_type and 'social' in img['type'].lower()):
                    used_types.add(img['filename'])
        
        return [img for img in image_files if img['filename'] not in used_types]
    
    def create_download_package(self, session_id: str) -> Dict:
        """Create downloadable Word document with content and embedded images"""
        return self.create_word_document(session_id)
    
    def register_image_file(self, session_id: str, image_path: str, 
                          image_type: str, filename: str):
        """Register an image file with the session"""
        if session_id not in self.sessions:
            self.sessions[session_id] = {'content_files': [], 'image_files': []}
        
        self.sessions[session_id]['image_files'].append({
            'filename': filename,
            'path': image_path,
            'type': image_type,
            'created': datetime.now().isoformat()
        })

# ADK Tool Functions

def generate_marketing_copy_tool(
    campaign_type: str,
    product_or_service: str,
    target_audience: str,
    key_message: str,
    call_to_action: str,
    tone: str = "professional",
    length: str = "medium"
) -> str:
    """
    ADK Tool: Generate marketing copy for campaigns
    
    Args:
        campaign_type: Type of campaign ('email', 'social', 'ad', 'landing_page', 'newsletter')
        product_or_service: What you're marketing
        target_audience: Who you're targeting
        key_message: Main message to convey
        call_to_action: Desired action from audience
        tone: Writing tone ('professional', 'casual', 'urgent', 'friendly', 'authoritative')
        length: Content length ('short', 'medium', 'long')
        
    Returns:
        JSON string with generated marketing copy
    """
    try:
        from vertexai.generative_models import GenerativeModel
        
        # Create marketing copy generation prompt
        prompt = f"""
You are a professional marketing copywriter. Generate compelling marketing copy based on these specifications:

CAMPAIGN DETAILS:
- Campaign Type: {campaign_type}
- Product/Service: {product_or_service}
- Target Audience: {target_audience}
- Key Message: {key_message}
- Call to Action: {call_to_action}
- Tone: {tone}
- Length: {length}

REQUIREMENTS:
- Create engaging, persuasive copy that resonates with the target audience
- Include compelling headlines and subheadings
- Incorporate the key message naturally
- End with a strong call-to-action
- Match the specified tone and length
- Format appropriately for the campaign type

Generate professional marketing copy that drives engagement and conversions.
"""
        
        # Generate copy using Gemini
        model = GenerativeModel("gemini-2.0-flash")
        response = model.generate_content(prompt)
        generated_copy = response.text
        
        return json.dumps({
            "success": True,
            "marketing_copy": generated_copy,
            "campaign_details": {
                "campaign_type": campaign_type,
                "product_or_service": product_or_service,
                "target_audience": target_audience,
                "key_message": key_message,
                "call_to_action": call_to_action,
                "tone": tone,
                "length": length
            },
            "word_count": len(generated_copy.split()),
            "character_count": len(generated_copy),
            "message": "Marketing copy generated successfully. Use save_marketing_content_tool to save it."
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        }, indent=2)

def save_marketing_content_tool(
    content: str,
    content_type: str,
    session_id: str,
    filename: Optional[str] = None
) -> str:
    """
    ADK Tool: Save marketing content to file
    
    Args:
        content: The marketing content to save
        content_type: Type of content ('email_copy', 'ad_copy', 'social_copy', 'landing_page', etc.)
        session_id: User session identifier
        filename: Optional custom filename
        
    Returns:
        JSON string with save results
    """
    try:
        initialize_content_manager()
        
        result = _content_manager.save_marketing_content(
            content=content,
            content_type=content_type,
            session_id=session_id,
            filename=filename
        )
        
        if result['success']:
            return json.dumps({
                "success": True,
                "message": f"Marketing content saved successfully as {result['filename']}",
                "file_details": {
                    "filename": result['filename'],
                    "path": result['path'],
                    "size_kb": result['size_kb'],
                    "content_type": content_type
                },
                "next_step": "Use create_download_package_tool to create a downloadable Word document"
            }, indent=2)
        else:
            return json.dumps({
                "success": False,
                "error": result['error']
            }, indent=2)
            
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        }, indent=2)

def create_download_package_tool(session_id: str) -> str:
    """
    ADK Tool: Create downloadable Word document with marketing content and embedded images
    
    Args:
        session_id: User session identifier
        
    Returns:
        JSON string with download document details
    """
    try:
        initialize_content_manager()
        
        result = _content_manager.create_download_package(session_id)
        
        if result['success']:
            # Create a file URL that works with ADK web interface
            file_url = f"/{result['download_url']}"
            
            return json.dumps({
                "success": True,
                "message": "Marketing campaign document created successfully!",
                "download_details": {
                    "document_filename": result['doc_filename'],
                    "size_mb": result['size_mb'],
                    "content_count": result['content_count'],
                    "image_count": result['image_count'],
                    "file_path": result['doc_path']
                },
                "download_info": {
                    "format": "Microsoft Word Document (.docx)",
                    "contains": "Marketing copy with embedded images ready for use",
                    "usage": "Open in Microsoft Word, Google Docs, or any Word processor"
                },
                "instructions": {
                    "step_1": "Click the download link below to save the document",
                    "step_2": "Open the document in Microsoft Word or Google Docs", 
                    "step_3": "Copy content and images directly into your email platform",
                    "step_4": "Customize as needed for your brand"
                },
                "download_link": {
                    "text": f"📄 Download Marketing Campaign Document",
                    "url": file_url,
                    "filename": result['doc_filename']
                }
            }, indent=2)
        else:
            return json.dumps({
                "success": False,
                "error": result['error']
            }, indent=2)
            
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        }, indent=2)

def register_session_image_tool(
    session_id: str,
    image_path: str,
    image_type: str,
    filename: str
) -> str:
    """
    ADK Tool: Register an image file with the user session
    
    Args:
        session_id: User session identifier  
        image_path: Full path to the image file
        image_type: Type of image generated
        filename: Image filename
        
    Returns:
        JSON string confirming registration
    """
    try:
        initialize_content_manager()
        
        _content_manager.register_image_file(
            session_id=session_id,
            image_path=image_path,
            image_type=image_type,
            filename=filename
        )
        
        return json.dumps({
            "success": True,
            "message": f"Image {filename} registered with session {session_id}",
            "image_details": {
                "filename": filename,
                "path": image_path,
                "type": image_type
            }
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": str(e)
        }, indent=2)

def list_session_content_tool(session_id: str) -> str:
    """
    ADK Tool: List all content created in a session
    
    Args:
        session_id: User session identifier
        
    Returns:
        JSON string with session content summary
    """
    try:
        initialize_content_manager()
        
        if session_id not in _content_manager.sessions:
            return json.dumps({
                "success": True,
                "message": "No content found for this session yet",
                "content_summary": {
                    "content_files": [],
                    "image_files": [],
                    "total_items": 0
                }
            }, indent=2)
        
        session_data = _content_manager.sessions[session_id]
        content_files = session_data.get('content_files', [])
        image_files = session_data.get('image_files', [])
        
        return json.dumps({
            "success": True,
            "message": f"Found {len(content_files)} content files and {len(image_files)} images",
            "content_summary": {
                "content_files": content_files,
                "image_files": image_files,
                "total_items": len(content_files) + len(image_files)
            },
            "next_actions": [
                "Generate more content with generate_marketing_copy_tool",
                "Create images with generate_marketing_image_tool", 
                "Create Word document with create_download_package_tool"
            ]
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            "success": False,
            "error": str(e)
        }, indent=2)